import io
import json
import os
import re
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st
from docx import Document
from openai import OpenAI
from pypdf import PdfReader

APP_DIR = Path(__file__).parent
DB_PATH = APP_DIR / "hvacr_reviews.db"

st.set_page_config(page_title="TechCheck HVAC&R", page_icon="⚙️", layout="wide", initial_sidebar_state="expanded")

THEME_CSS = r"""
<style>
:root {
  --navy:#12233f; --navy2:#182d4d; --orange:#f28c28; --orange2:#ffab4a;
  --ink:#172033; --muted:#687386; --line:#e3e8ef; --panel:#ffffff; --bg:#f4f6f9;
}
.stApp { background: var(--bg); color: var(--ink); }
[data-testid="stHeader"] { background: transparent; }
[data-testid="stToolbar"], footer { visibility: hidden; }
.block-container { padding-top: 2rem; padding-bottom: 3rem; max-width: 1500px; }
[data-testid="stSidebar"] { background: linear-gradient(180deg,var(--navy) 0%,#0d1a30 100%); border-right:0; }
[data-testid="stSidebar"] * { color:#eef4fb; }
[data-testid="stSidebar"] .stRadio label { padding:.45rem .6rem; border-radius:9px; margin:.15rem 0; }
[data-testid="stSidebar"] .stRadio label:hover { background:rgba(255,255,255,.08); }
[data-testid="stSidebar"] [data-baseweb="radio"] > div:first-child { border-color:#9eb0c9; }
.brand { padding:.35rem .2rem 1.2rem; border-bottom:1px solid rgba(255,255,255,.13); margin-bottom:1rem; }
.brand-title { font-size:1.35rem; font-weight:800; letter-spacing:.2px; color:#fff; }
.brand-title span { color:var(--orange2); }
.brand-sub { font-size:.78rem; color:#b8c6d9; margin-top:.22rem; }
.page-kicker { color:var(--orange); font-weight:800; text-transform:uppercase; letter-spacing:.11em; font-size:.74rem; margin-bottom:.25rem; }
.page-title { color:var(--navy); font-weight:850; font-size:2.2rem; line-height:1.12; margin:0; }
.page-subtitle { color:var(--muted); margin:.45rem 0 1.35rem; font-size:.98rem; }
[data-testid="stMetric"] { background:var(--panel); border:1px solid var(--line); padding:1rem 1.1rem; border-radius:14px; box-shadow:0 7px 22px rgba(18,35,63,.055); }
[data-testid="stMetricLabel"] { color:var(--muted); }
[data-testid="stMetricValue"] { color:var(--navy); font-weight:800; }
.stButton > button, .stDownloadButton > button { border-radius:9px; font-weight:700; border:1px solid #cfd7e3; min-height:2.6rem; }
.stButton > button[kind="primary"] { background:var(--orange); border-color:var(--orange); color:#fff; }
.stButton > button[kind="primary"]:hover { background:#dc7719; border-color:#dc7719; color:#fff; }
.stTextInput input, .stNumberInput input, .stTextArea textarea, [data-baseweb="select"] > div { border-radius:9px !important; background:#fff !important; }
[data-testid="stFileUploader"] { background:#fff; border:1px dashed #b8c3d2; border-radius:12px; padding:.45rem; }
[data-testid="stDataFrame"], [data-testid="stTable"] { border:1px solid var(--line); border-radius:12px; overflow:hidden; background:#fff; }
[data-testid="stAlert"] { border-radius:11px; }
.stTabs [data-baseweb="tab-list"] { gap:.35rem; }
.stTabs [data-baseweb="tab"] { border-radius:8px 8px 0 0; padding:.7rem 1rem; }
.stTabs [aria-selected="true"] { color:var(--orange) !important; border-bottom-color:var(--orange) !important; }
[data-testid="stChatMessage"] { background:#fff; border:1px solid var(--line); border-radius:14px; padding:.3rem .8rem; margin:.5rem 0; box-shadow:0 5px 18px rgba(18,35,63,.04); }
[data-testid="stChatInput"] { border-radius:12px; }
.section-card { background:#fff; border:1px solid var(--line); border-radius:14px; padding:1rem 1.1rem; box-shadow:0 7px 22px rgba(18,35,63,.045); margin-bottom:1rem; }
.sidebar-note { color:#b8c6d9 !important; font-size:.73rem; line-height:1.45; padding-top:1rem; }
hr { border-color:var(--line); }
</style>
"""
st.markdown(THEME_CSS, unsafe_allow_html=True)

SYSTEM_PROMPT = r"""
You are an independent Senior HVAC&R Technical Manager reviewing technician reports before they are issued to a customer.
Protect both contractor and customer. Do not merely agree with the technician. Require evidence, identify assumptions,
challenge expensive repairs, check labour reasonableness, and distinguish symptoms from root cause.

ADAPTIVE REVIEW
First classify the job and activate only relevant modules: refrigeration/cold room, DX air conditioning, VRF/VRV,
controls/BMS, electrical, hydronics, chiller, cooling tower, heat pump/pool heating, ventilation/exhaust, PM audit,
quotation/additional works, compliance, or general mechanical.

MANDATORY REVIEW
1. Job summary: complaint, findings, root cause, work completed, outstanding work, criticality.
2. Diagnosis: Supported / Possibly Supported / Not Supported, with evidence for and against.
3. Activated modules and why.
4. Technical analysis of all readings and observations. Never invent missing readings.
5. Missing tests/evidence, prioritised by what most increases diagnostic confidence.
6. Parts review: each recommended/replaced part, justification YES/NO/UNSURE, confidence 0-100%, alternatives.
7. Labour analysis: estimate reasonable onsite labour range, compare with claimed hours, and account for access,
   defrosting, waiting/stabilisation, multiple technicians, commissioning, documentation and return visits.
8. Report quality: clarity, sequence, professionalism, customer defensibility, photos/evidence.
9. Risks: callback, safety, financial, warranty, compliance, reputation.
10. Senior Technician Challenge: alternatives, cheapest sensible next step, strongest missing test, symptom vs root cause,
    and whether the recommendation is sufficiently evidenced to spend the customer's money.
11. Coaching feedback: what was done well and what should improve.
12. Final decision: Approve / Approve with Comments / Request Further Information / Reject Recommendation.

COMPANY RULES
- Expensive component replacements require direct evidence, not inference alone.
- Compressor recommendations should normally include electrical checks, operating readings, system condition and cause of failure.
- TXV/EEV recommendations should normally address load, airflow, charge, restriction, bulb mounting/insulation, equaliser and controls.
- Controller/sensor faults require calibration, offsets, placement, setpoint/deadband and actual-temperature comparison.
- Refrigeration conclusions should use refrigerant type and saturated temperatures where possible.
- A successful outcome does not retroactively prove every earlier diagnosis.
- Avoid false precision. State uncertainty clearly.

Return valid JSON only using this exact structure:
{
  "job_classification": {"type": "", "activated_modules": [""], "criticality": "Low|Medium|High|Critical"},
  "executive_summary": "",
  "diagnosis": {"rating": "Supported|Possibly Supported|Not Supported", "confidence": 0, "analysis": "", "root_cause": ""},
  "technical_findings": [{"finding": "", "assessment": "Good|Concern|Missing|Neutral", "explanation": ""}],
  "missing_evidence": [{"priority": "High|Medium|Low", "item": "", "why": ""}],
  "parts_review": [{"part": "", "justified": "YES|NO|UNSURE", "confidence": 0, "reason": "", "alternatives": [""]}],
  "labour_review": {"claimed_hours": 0, "reasonable_low": 0, "reasonable_high": 0, "rating": "Unusually Low|Reasonable|High|Excessive|Cannot Assess", "analysis": ""},
  "risk_review": [{"risk": "Callback|Safety|Financial|Warranty|Compliance|Reputation", "level": "Low|Medium|High", "explanation": ""}],
  "senior_technician_challenge": [{"question": "", "answer": ""}],
  "questions_for_technician": [""],
  "coaching": {"done_well": [""], "improvements": [""], "better_approach": [""]},
  "scores": {"fault_finding": 0, "evidence": 0, "technical_quality": 0, "efficiency": 0, "labour_accuracy": 0, "report_quality": 0, "customer_confidence": 0, "overall_100": 0},
  "final_decision": {"decision": "Approve|Approve with Comments|Request Further Information|Reject Recommendation", "confidence": 0, "justification": ""},
  "customer_ready_summary": ""
}
"""



def page_header(title: str, subtitle: str = "", kicker: str = "TechCheck HVAC&R") -> None:
    st.markdown(
        f'<div class="page-kicker">{kicker}</div><h1 class="page-title">{title}</h1>'
        + (f'<div class="page-subtitle">{subtitle}</div>' if subtitle else '<div style="height:.8rem"></div>'),
        unsafe_allow_html=True,
    )


def init_db() -> None:
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("""
        CREATE TABLE IF NOT EXISTS reviews (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT NOT NULL,
            work_order TEXT,
            customer TEXT,
            site TEXT,
            technician TEXT,
            report_name TEXT,
            claimed_hours REAL,
            mode TEXT,
            decision TEXT,
            overall_score INTEGER,
            report_text TEXT,
            review_json TEXT
        )
        """)
        conn.commit()


def save_review(meta: dict[str, Any], report_text: str, result: dict[str, Any], mode: str) -> int:
    decision = result.get("final_decision", {}).get("decision", "")
    score = int(result.get("scores", {}).get("overall_100", 0) or 0)
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.execute(
            """INSERT INTO reviews
            (created_at, work_order, customer, site, technician, report_name, claimed_hours, mode,
             decision, overall_score, report_text, review_json)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                datetime.now().isoformat(timespec="seconds"), meta.get("work_order", ""), meta.get("customer", ""),
                meta.get("site", ""), meta.get("technician", ""), meta.get("report_name", ""),
                float(meta.get("hours", 0)), mode, decision, score, report_text, json.dumps(result),
            ),
        )
        conn.commit()
        return int(cur.lastrowid)


def load_reviews() -> pd.DataFrame:
    with sqlite3.connect(DB_PATH) as conn:
        return pd.read_sql_query(
            "SELECT id, created_at, work_order, customer, site, technician, report_name, claimed_hours, mode, decision, overall_score FROM reviews ORDER BY id DESC",
            conn,
        )


def get_review(review_id: int) -> tuple[str, dict[str, Any]]:
    with sqlite3.connect(DB_PATH) as conn:
        row = conn.execute("SELECT report_text, review_json FROM reviews WHERE id = ?", (review_id,)).fetchone()
    if not row:
        raise ValueError("Review not found")
    return row[0], json.loads(row[1])


def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(lines)
    if name.endswith((".txt", ".csv", ".md")):
        return data.decode("utf-8", errors="replace")
    raise ValueError("Supported files: PDF, DOCX, TXT, CSV and MD.")


def parse_json_response(text: str) -> dict[str, Any]:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    start, end = text.find("{"), text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("The AI response did not contain JSON.")
    return json.loads(text[start:end + 1])


def live_review(report: str, hours: float, context: str, standards: str, model: str) -> dict[str, Any]:
    client = OpenAI(api_key=st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY")))
    user_input = f"""Review this HVAC&R technician report.
Claimed total labour hours: {hours:.2f}
Additional company/job context: {context or 'None provided'}
Company-specific standards: {standards or 'Use the default standards in the system instructions.'}

TECHNICIAN REPORT:
{report}
"""
    response = client.responses.create(model=model, instructions=SYSTEM_PROMPT, input=user_input)
    return parse_json_response(response.output_text)



def technical_chat(question: str, report: str, review: dict[str, Any] | None, model: str) -> str:
    api_key = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
    if not api_key:
        raise ValueError("Live AI is not configured. Add OPENAI_API_KEY in Streamlit Secrets.")
    client = OpenAI(api_key=api_key)
    review_context = json.dumps(review, indent=2) if review else "No structured review is available."
    prompt = f"""You are the Technical Leading Hand: a senior HVAC&R service manager helping a supervisor or contractor analyse a job.

Answer the user's question directly and practically. Base your answer on the technician report and structured review supplied below.
Do not invent readings, standards, manufacturer instructions, fault codes or events. Clearly separate confirmed facts, likely interpretations and missing evidence.
Challenge unsupported component replacement and unreasonable labour, but acknowledge justified work.
Where useful, provide the technical reasoning, what to ask the technician, the next diagnostic test, and whether the report should be amended.
This is advisory assistance only; do not claim to certify safety, compliance or workmanship.

USER QUESTION:
{question}

TECHNICIAN REPORT:
{report or 'No technician report supplied.'}

STRUCTURED REVIEW:
{review_context}
"""
    response = client.responses.create(model=model, input=prompt)
    return response.output_text.strip()

def demo_review(report: str, hours: float) -> dict[str, Any]:
    lower = report.lower()
    modules = []
    for keywords, module in [
        (("cool room", "freezer", "txv", "r404", "evaporator"), "Refrigeration / Cold Room"),
        (("controller", "sensor", "offset", "thermostat"), "Controls"),
        (("chiller",), "Chiller / Hydronics"),
        (("vrf", "vrv", "refnet"), "VRF / VRV"),
        (("pm", "preventative maintenance"), "PM Audit"),
        (("quote", "quotation", "additional works"), "Quotation / Additional Works"),
    ]:
        if any(k in lower for k in keywords):
            modules.append(module)
    modules = modules or ["General HVAC&R"]
    offset, txv, compressor = "offset" in lower, "txv" in lower, "compressor" in lower
    heavy_defrost = "ice" in lower and any(k in lower for k in ("melt", "defrost", "block"))
    reasonable_low, reasonable_high = (8.0, 14.0) if heavy_defrost else (2.0, 7.0)
    labour_rating = "Reasonable" if reasonable_low <= hours <= reasonable_high + 1 else ("High" if hours > reasonable_high else "Unusually Low")
    key_part = "Compressor" if compressor else ("TXV" if txv else "No major part identified")
    return {
        "job_classification": {"type": modules[0], "activated_modules": modules, "criticality": "Medium"},
        "executive_summary": "Demonstration review only. It uses basic rules to prove the workflow; select Live AI for a report-specific technical assessment.",
        "diagnosis": {"rating": "Possibly Supported", "confidence": 65, "analysis": "The repair path may be plausible, but demonstration mode cannot perform full technical reasoning.", "root_cause": "Controller configuration/calibration appears relevant." if offset else "Insufficient information to confirm a single root cause."},
        "technical_findings": [
            {"finding": "Adaptive modules", "assessment": "Good", "explanation": ", ".join(modules)},
            {"finding": "Controller calibration", "assessment": "Good" if offset else "Missing", "explanation": "Offsets and actual-temperature comparison are key control checks."},
            {"finding": f"{key_part} justification", "assessment": "Concern" if (txv or compressor) else "Neutral", "explanation": "Major components require an evidence chain, not symptoms alone."},
        ],
        "missing_evidence": [
            {"priority": "High", "item": "Evidence chain", "why": "Link each conclusion and part recommendation to the tests that support it."},
            {"priority": "Medium", "item": "Time breakdown", "why": "Break labour down by visit and task so the charge is defensible."},
        ],
        "parts_review": [{"part": key_part, "justified": "UNSURE", "confidence": 50, "reason": "Demonstration mode cannot validate component failure.", "alternatives": ["Controls", "Sensor calibration", "Airflow/load", "Charge or restriction"]}],
        "labour_review": {"claimed_hours": hours, "reasonable_low": reasonable_low, "reasonable_high": reasonable_high, "rating": labour_rating, "analysis": "Assess defrost time, access, stabilisation, return visits and documentation."},
        "risk_review": [
            {"risk": "Callback", "level": "Medium", "explanation": "Root cause and contributing faults must be separated."},
            {"risk": "Financial", "level": "Medium", "explanation": "Weak part justification can lead to customer challenge."},
        ],
        "senior_technician_challenge": [
            {"question": "What else could cause the symptoms?", "answer": "Controls, sensor placement, airflow/load, charge or restriction."},
            {"question": "What test most increases confidence?", "answer": "Repeatable measurements before and after each isolated correction."},
            {"question": "Would you spend your own money?", "answer": "Only after the evidence supports the recommendation."},
        ],
        "questions_for_technician": ["Provide labour by visit/task.", "Confirm the primary root cause.", "List tests that ruled out alternatives."],
        "coaching": {"done_well": ["Recorded observations."], "improvements": ["Separate hypothesis from confirmed diagnosis.", "Link labour to activities."], "better_approach": ["Correct control and installation defects first, stabilise, then reassess."]},
        "scores": {"fault_finding": 6, "evidence": 5, "technical_quality": 6, "efficiency": 6, "labour_accuracy": 6, "report_quality": 6, "customer_confidence": 6, "overall_100": 58},
        "final_decision": {"decision": "Request Further Information", "confidence": 65, "justification": "Use Live AI before relying on this result for a real commercial decision."},
        "customer_ready_summary": "Demonstration result only. A report-specific live AI review is required before issue.",
    }


def render_review(r: dict[str, Any], key_prefix: str = "current") -> None:
    decision = r.get("final_decision", {})
    st.subheader(f"Decision: {decision.get('decision', '—')}")
    st.write(decision.get("justification", ""))
    scores = r.get("scores", {})
    cols = st.columns(4)
    cols[0].metric("Overall", f"{scores.get('overall_100', 0)}/100")
    cols[1].metric("Diagnosis confidence", f"{r.get('diagnosis', {}).get('confidence', 0)}%")
    cols[2].metric("Labour", r.get("labour_review", {}).get("rating", "—"))
    cols[3].metric("Criticality", r.get("job_classification", {}).get("criticality", "—"))

    tabs = st.tabs(["Summary", "Technical", "Labour & Parts", "Risk & Challenge", "Coaching", "Customer Copy", "Raw JSON"])
    with tabs[0]:
        st.markdown("### Classification")
        st.write(r.get("job_classification", {}))
        st.markdown("### Executive summary")
        st.write(r.get("executive_summary", ""))
        st.markdown("### Diagnosis")
        st.write(r.get("diagnosis", {}))
    with tabs[1]:
        for item in r.get("technical_findings", []):
            st.markdown(f"**{item.get('assessment', '')}: {item.get('finding', '')}**")
            st.write(item.get("explanation", ""))
        st.markdown("### Missing evidence")
        st.dataframe(r.get("missing_evidence", []), use_container_width=True)
        st.markdown("### Questions for technician")
        for q in r.get("questions_for_technician", []):
            st.write(f"• {q}")
    with tabs[2]:
        st.markdown("### Labour")
        st.write(r.get("labour_review", {}))
        st.markdown("### Parts")
        st.dataframe(r.get("parts_review", []), use_container_width=True)
    with tabs[3]:
        st.markdown("### Risks")
        st.dataframe(r.get("risk_review", []), use_container_width=True)
        st.markdown("### Senior Technician Challenge")
        for item in r.get("senior_technician_challenge", []):
            st.markdown(f"**{item.get('question', '')}**")
            st.write(item.get("answer", ""))
    with tabs[4]:
        coaching = r.get("coaching", {})
        for heading, key in [("Done well", "done_well"), ("Improvements", "improvements"), ("Better approach", "better_approach")]:
            st.markdown(f"### {heading}")
            for x in coaching.get(key, []):
                st.write(f"• {x}")
    with tabs[5]:
        st.text_area("Customer-ready summary", r.get("customer_ready_summary", ""), height=240, key=f"{key_prefix}_customer")
    with tabs[6]:
        st.json(r)
    st.download_button("Download review JSON", json.dumps(r, indent=2), file_name=f"hvacr_review_{datetime.now():%Y%m%d_%H%M}.json", mime="application/json", key=f"{key_prefix}_download")




QUOTE_REVIEW_PROMPT = r"""
You are an independent Senior HVAC&R Service Manager and commercial estimator reviewing an HVAC&R quotation before issue or approval.

Your job is to determine whether the scope, labour allowance, material/equipment allowance and commercial assumptions are reasonable for the work described. Think like an experienced HVAC&R contractor AND an experienced client-side contract supervisor.

Do not merely accept the entered hours or dollar allowance. Challenge them against the stated scope and normal HVAC&R work practices. However, do not invent exact market prices or claim a benchmark is certain when the quote lacks information. Use ranges and explain uncertainty.

For every quote:
- Identify exactly what work is being proposed.
- Decide whether the scope is clear enough to price and approve.
- Estimate a reasonable onsite labour range based on the stated work only. Account for likely crew size, isolation, access, removal, installation, testing, commissioning, cleanup and documentation.
- Compare the entered labour hours with your reasonable range and rate it Unusually Low / Reasonable / High / Excessive / Cannot Assess.
- Assess whether the materials/equipment allowance appears plausible for the stated scope. If it appears unusually high or low, say so, but do not fabricate supplier pricing.
- Identify missing scope items, exclusions, assumptions, access requirements, permits, after-hours work, EWP/scaffold/crane, builder's works, electrical, controls, condensate, refrigeration, recovery/recharge, commissioning, disposal and warranty where relevant.
- Identify technical or commercial risks that could lead to a variation, callback or dispute.
- Ask the specific questions that should be answered before approval.
- Give a final recommendation: Approve / Approve with Comments / Request Further Information / Reject / Reprice.

Important HVAC&R reasoning rules:
- A simple condensate drain repair/replacement should not automatically attract major labour or material allowances unless the scope explains difficult access, long runs, pumps, builder's works, ceiling works, EWP/scaffold, after-hours work or multiple technicians.
- Compressor, coil, chiller, VRF and major refrigeration work require broader allowances and commissioning steps.
- Refrigerant work should address recovery, leak testing/pressure testing, evacuation, recharge and commissioning where applicable.
- Labour is labour-hours, not elapsed job duration. If two technicians attend for 5 hours, that is 10 labour-hours.
- When the description is too vague to benchmark confidently, say Cannot Assess and identify what information is missing.

Return valid JSON only using this exact structure:
{
  "commercial_risk": "Low|Medium|High",
  "scope_summary": "",
  "scope_clarity": "Clear|Partly Clear|Insufficient",
  "labour_review": {
    "claimed_hours": 0,
    "reasonable_low": 0,
    "reasonable_high": 0,
    "rating": "Unusually Low|Reasonable|High|Excessive|Cannot Assess",
    "analysis": ""
  },
  "materials_review": {
    "claimed_allowance": 0,
    "rating": "Low|Plausible|High|Excessive|Cannot Assess",
    "analysis": ""
  },
  "items_to_check": [{"priority":"High|Medium|Low","item":"","reason":""}],
  "commercial_risks": [{"risk":"","level":"Low|Medium|High","reason":""}],
  "questions_before_approval": [""],
  "final_recommendation": {
    "decision": "Approve|Approve with Comments|Request Further Information|Reject|Reprice",
    "confidence": 0,
    "justification": ""
  }
}
"""


def live_quote_review(notes: str, labour_hours: float, material_cost: float, quote_ref: str, model: str) -> dict[str, Any]:
    api_key = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
    if not api_key:
        raise ValueError("Live AI is not configured. Add OPENAI_API_KEY in Streamlit Secrets.")
    client = OpenAI(api_key=api_key)
    prompt = f"""Review this HVAC&R quotation.
Quote reference: {quote_ref or 'Not provided'}
Claimed labour hours: {labour_hours:.2f}
Materials and equipment allowance: ${material_cost:,.2f}

SCOPE / QUOTATION DETAILS:
{notes}
"""
    response = client.responses.create(model=model, instructions=QUOTE_REVIEW_PROMPT, input=prompt)
    return parse_json_response(response.output_text)

def estimate_quote_risk(notes: str, labour_hours: float, material_cost: float) -> dict[str, Any]:
    text = notes.lower()
    flags: list[str] = []
    if labour_hours <= 0:
        flags.append("No labour allowance entered.")
    if material_cost <= 0:
        flags.append("No material allowance entered.")
    if any(x in text for x in ["compressor", "chiller", "coil replacement", "major leak"]):
        flags.append("Major component work should include direct diagnostic evidence and cause-of-failure analysis.")
    if "refrigerant" in text and not any(x in text for x in ["kg", "charge", "recover", "recovery"]):
        flags.append("Refrigerant quantity and recovery/recharge method are unclear.")
    if any(x in text for x in ["access", "ewp", "scissor", "crane", "after hours"]):
        flags.append("Confirm access, hire and after-hours allowances are separately identified.")
    risk = "Low" if len(flags) <= 1 else "Medium" if len(flags) <= 3 else "High"
    return {"risk": risk, "flags": flags}


def technician_summary(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty or "technician" not in df.columns:
        return pd.DataFrame()
    data = df.copy()
    data["technician"] = data["technician"].fillna("").replace("", "Unassigned")
    grouped = data.groupby("technician", dropna=False).agg(
        reviews=("id", "count"),
        average_score=("overall_score", "mean"),
        hours_reviewed=("claimed_hours", "sum"),
    ).reset_index()
    grouped["average_score"] = grouped["average_score"].round(0).astype(int)
    flagged = data[data["decision"].isin(["Request Further Information", "Reject Recommendation"])].groupby("technician").size()
    grouped["flagged"] = grouped["technician"].map(flagged).fillna(0).astype(int)
    return grouped.sort_values(["average_score", "reviews"], ascending=[False, False])

def require_access() -> None:
    expected = st.secrets.get("APP_PASSWORD", os.getenv("APP_PASSWORD", ""))
    if not expected:
        return
    if st.session_state.get("authenticated"):
        return
    st.title("TechCheck HVAC&R")
    st.caption("Private beta access")
    password = st.text_input("Access password", type="password")
    if st.button("Sign in", type="primary"):
        if password == expected:
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("Incorrect password")
    st.stop()


init_db()
require_access()

st.sidebar.markdown(
    '<div class="brand"><div class="brand-title">TechCheck <span>HVAC&R</span></div>'
    '<div class="brand-sub">Technical QA & diagnostic support</div></div>',
    unsafe_allow_html=True,
)
page = st.sidebar.radio(
    "Navigation",
    ["Dashboard", "Review report", "Sounding Board", "Quote review", "Technicians", "Analytics", "History", "Company standards", "Settings"],
    label_visibility="collapsed",
)
st.sidebar.markdown('<div class="sidebar-note">AI supports supervisor judgement and does not certify safety, compliance or workmanship.</div>', unsafe_allow_html=True)

if page == "Dashboard":
    page_header("Technical QA Dashboard", "Review performance, flagged jobs and recent technical QA activity.", "Overview")
    reviews = load_reviews()
    if reviews.empty:
        st.info("No saved reviews yet. Open ‘Review report’ to create the first one.")
    else:
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Reviews", len(reviews))
        c2.metric("Average score", f"{reviews['overall_score'].mean():.0f}/100")
        c3.metric("Hours reviewed", f"{reviews['claimed_hours'].sum():.1f}")
        flagged = reviews[reviews["decision"].isin(["Request Further Information", "Reject Recommendation"])]
        c4.metric("Flagged", len(flagged))
        st.markdown("## Recent reviews")
        st.dataframe(reviews.head(10), use_container_width=True, hide_index=True)
        st.markdown("## Decision mix")
        st.bar_chart(reviews["decision"].value_counts())

elif page == "Review report":
    page_header("Review technician report", "Upload service notes and receive a structured technical, labour and risk assessment.", "Technical review")
    live_key = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
    if not live_key:
        st.warning("Live AI is not configured. Demonstration mode uses generic rules and can produce similar results for similar reports.")
    with st.sidebar:
        st.header("Review settings")
        mode = st.radio("Review mode", ["Live AI", "Demonstration"], index=0 if live_key else 1)
        model = st.text_input("Model", value=st.secrets.get("OPENAI_MODEL", os.getenv("OPENAI_MODEL", "gpt-4.1-mini")))
        hours = st.number_input("Claimed total hours", min_value=0.0, value=0.0, step=0.25)
        context = st.text_area("Job context", placeholder="Access, number of technicians, SLA, travel policy, operating conditions...")
    cols = st.columns(2)
    with cols[0]:
        work_order = st.text_input("Work order / job number")
        customer = st.text_input("Customer")
        site = st.text_input("Site")
    with cols[1]:
        technician = st.text_input("Technician")
        uploaded = st.file_uploader("Upload report", type=["pdf", "docx", "txt", "csv", "md"])
    report_text = st.text_area("Technician notes", height=340, placeholder="Paste the complete job notes here...")
    report_name = "Pasted notes"
    if uploaded:
        report_name = uploaded.name
        try:
            extracted = extract_text(uploaded)
            if extracted.strip():
                report_text = extracted
                st.success(f"Extracted {len(extracted):,} characters from {uploaded.name}.")
            else:
                st.warning("No selectable text was found. A scanned PDF requires OCR or image-capable processing.")
        except Exception as exc:
            st.error(str(exc))
    standards = st.session_state.get("company_standards", st.secrets.get("COMPANY_STANDARDS", ""))
    if st.button("Run technical review", type="primary", use_container_width=True):
        if not report_text.strip():
            st.error("Paste notes or upload a report first.")
        elif mode == "Live AI" and not live_key:
            st.error("Add OPENAI_API_KEY in Streamlit Secrets first.")
        else:
            with st.spinner("Reviewing report..."):
                try:
                    result = live_review(report_text, hours, context, standards, model) if mode == "Live AI" else demo_review(report_text, hours)
                    meta = {"work_order": work_order, "customer": customer, "site": site, "technician": technician, "report_name": report_name, "hours": hours}
                    review_id = save_review(meta, report_text, result, mode)
                    st.session_state["review"] = result
                    st.session_state["review_id"] = review_id
                    st.success(f"Review saved as #{review_id}.")
                except Exception as exc:
                    st.exception(exc)
    if "review" in st.session_state:
        render_review(st.session_state["review"])

elif page == "Sounding Board":
    page_header("Sounding Board", "", "Diagnostic support")

    live_key = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
    model = st.secrets.get("OPENAI_MODEL", os.getenv("OPENAI_MODEL", "gpt-4.1-mini"))

    if "chat_messages" not in st.session_state:
        st.session_state["chat_messages"] = [
            {
                "role": "assistant",
                "content": "Describe the symptoms, operating conditions and any readings you have. I’ll help narrow the likely causes and prioritise the next checks."
            }
        ]

    for message in st.session_state["chat_messages"]:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    question = st.chat_input("Describe the symptoms...")
    if question:
        st.session_state["chat_messages"].append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)
        with st.chat_message("assistant"):
            if not live_key:
                answer = "Live AI is not configured yet. Add OPENAI_API_KEY in Streamlit App settings → Secrets, then reboot the app."
                st.error(answer)
            else:
                with st.spinner("Thinking through the fault..."):
                    try:
                        conversation = "\n\n".join(
                            f"{m['role'].upper()}: {m['content']}"
                            for m in st.session_state["chat_messages"][-12:]
                        )
                        answer = technical_chat(conversation, "", None, model)
                        st.markdown(answer)
                    except Exception as exc:
                        answer = f"Unable to answer: {exc}"
                        st.error(answer)
        st.session_state["chat_messages"].append({"role": "assistant", "content": answer})

    c1, c2 = st.columns(2)
    with c1:
        if st.button("Clear chat"):
            st.session_state["chat_messages"] = [
                {
                    "role": "assistant",
                    "content": "Describe the symptoms, operating conditions and any readings you have. I’ll help narrow the likely causes and prioritise the next checks."
                }
            ]
            st.rerun()
    with c2:
        if len(st.session_state["chat_messages"]) > 1:
            transcript = "\n\n".join(f"{m['role'].upper()}: {m['content']}" for m in st.session_state["chat_messages"])
            st.download_button("Download chat transcript", transcript, file_name="sounding_board_chat.txt", mime="text/plain")


elif page == "Quote review":
    page_header("Review quotation", "Check scope clarity, labour allowances, exclusions and commercial risk before issue or approval.", "Commercial QA")
    q1, q2, q3 = st.columns(3)
    with q1:
        quote_ref = st.text_input("Quote reference")
    with q2:
        quote_labour = st.number_input("Labour hours", min_value=0.0, step=0.5)
    with q3:
        material_cost = st.number_input("Materials and equipment allowance ($)", min_value=0.0, step=50.0)
    quote_notes = st.text_area("Scope and quotation details", height=330, placeholder="Paste the proposed scope, labour, materials, access requirements, exclusions and warranty details...")
    if st.button("Review quotation", type="primary", use_container_width=True):
        if not quote_notes.strip():
            st.error("Paste the quotation scope first.")
        else:
            api_key = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
            if api_key:
                with st.spinner("Reviewing scope, labour and commercial risk..."):
                    try:
                        result = live_quote_review(quote_notes, quote_labour, material_cost, quote_ref, model)
                        st.session_state["quote_review"] = result
                        st.session_state["quote_review_mode"] = "Live AI"
                    except Exception as exc:
                        st.error(f"Live quote review failed: {exc}")
            else:
                result = estimate_quote_risk(quote_notes, quote_labour, material_cost)
                st.session_state["quote_review"] = result
                st.session_state["quote_review_mode"] = "Preliminary rules"
                st.warning("Live AI is not configured, so this is only a preliminary rule-based check.")
    if "quote_review" in st.session_state:
        qr = st.session_state["quote_review"]
        mode_used = st.session_state.get("quote_review_mode", "")
        if "commercial_risk" in qr:
            c1, c2, c3 = st.columns(3)
            c1.metric("Commercial risk", qr.get("commercial_risk", "—"))
            labour = qr.get("labour_review", {})
            materials = qr.get("materials_review", {})
            c2.metric("Labour assessment", labour.get("rating", "—"), f"{quote_labour:.1f} h claimed")
            c3.metric("Materials assessment", materials.get("rating", "—"), f"${material_cost:,.0f} claimed")

            st.markdown("### Review summary")
            st.write(qr.get("scope_summary", ""))
            st.caption(f"Scope clarity: {qr.get('scope_clarity', '—')} · Review mode: {mode_used}")

            left, right = st.columns(2)
            with left:
                st.markdown("### Labour")
                low = labour.get("reasonable_low", 0)
                high = labour.get("reasonable_high", 0)
                if low or high:
                    st.write(f"**Reasonable range:** {low:g}–{high:g} labour-hours")
                st.write(labour.get("analysis", ""))
            with right:
                st.markdown("### Materials & equipment")
                st.write(materials.get("analysis", ""))

            st.markdown("### Items to check")
            items = qr.get("items_to_check", [])
            if items:
                for item in items:
                    st.markdown(f"**{item.get('priority','')} — {item.get('item','')}**")
                    st.write(item.get("reason", ""))
            else:
                st.success("No material scope omissions were identified from the information supplied.")

            st.markdown("### Commercial risks")
            risks = qr.get("commercial_risks", [])
            if risks:
                st.dataframe(risks, use_container_width=True, hide_index=True)
            else:
                st.write("No significant commercial risks identified from the supplied scope.")

            st.markdown("### Questions before approval")
            questions = qr.get("questions_before_approval", [])
            if questions:
                for q in questions:
                    st.write(f"• {q}")
            else:
                st.write("No additional questions identified.")

            final = qr.get("final_recommendation", {})
            st.markdown("### Recommendation")
            st.subheader(final.get("decision", "—"))
            st.write(final.get("justification", ""))
            if final.get("confidence") is not None:
                st.caption(f"Confidence: {final.get('confidence', 0)}%")
        else:
            c1, c2, c3 = st.columns(3)
            c1.metric("Commercial risk", qr["risk"])
            c2.metric("Labour allowance", f"{quote_labour:.1f} h")
            c3.metric("Materials allowance", f"${material_cost:,.0f}")
            st.markdown("### Items to check")
            if qr["flags"]:
                for flag in qr["flags"]:
                    st.write(f"• {flag}")
            else:
                st.success("No obvious omissions were detected by the preliminary checks.")
            st.caption("Preliminary rule-based check only. Add OPENAI_API_KEY in Streamlit Secrets to enable backend AI reasoning.")

elif page == "Technicians":
    page_header("Technician performance", "Use review outcomes to identify coaching strengths, recurring gaps and report-quality trends.", "People")
    reviews = load_reviews()
    tech = technician_summary(reviews)
    if tech.empty:
        st.info("Technician analytics will appear after reviews are saved with technician names.")
    else:
        st.dataframe(tech, use_container_width=True, hide_index=True)
        selected_tech = st.selectbox("Technician detail", tech["technician"].tolist())
        subset = reviews[reviews["technician"].fillna("").replace("", "Unassigned") == selected_tech]
        c1, c2, c3 = st.columns(3)
        c1.metric("Reviews", len(subset))
        c2.metric("Average score", f"{subset['overall_score'].mean():.0f}/100")
        c3.metric("Flagged", len(subset[subset["decision"].isin(["Request Further Information", "Reject Recommendation"])]))
        st.dataframe(subset.head(20), use_container_width=True, hide_index=True)

elif page == "Analytics":
    page_header("Technical analytics", "See review volume, labour exposure, decisions and recurring quality signals.", "Insights")
    reviews = load_reviews()
    if reviews.empty:
        st.info("Analytics will populate after the first saved review.")
    else:
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Total reviews", len(reviews))
        c2.metric("Average score", f"{reviews['overall_score'].mean():.0f}/100")
        c3.metric("Labour reviewed", f"{reviews['claimed_hours'].sum():.1f} h")
        c4.metric("Unique technicians", reviews["technician"].replace("", pd.NA).dropna().nunique())
        left, right = st.columns(2)
        with left:
            st.markdown("### Decision distribution")
            st.bar_chart(reviews["decision"].value_counts())
        with right:
            st.markdown("### Average score by technician")
            tech = technician_summary(reviews)
            if not tech.empty:
                st.bar_chart(tech.set_index("technician")["average_score"])
        st.markdown("### Review register")
        st.dataframe(reviews, use_container_width=True, hide_index=True)

elif page == "History":
    page_header("Review history", "Open previous assessments and export the review register.", "Records")
    reviews = load_reviews()
    if reviews.empty:
        st.info("No reviews saved yet.")
    else:
        st.dataframe(reviews, use_container_width=True, hide_index=True)
        selected = st.selectbox("Open review", reviews["id"].tolist(), format_func=lambda x: f"#{x} — {reviews.loc[reviews['id'] == x, 'work_order'].iloc[0] or 'No work order'}")
        if selected:
            report, result = get_review(int(selected))
            with st.expander("Original technician report"):
                st.text(report)
            render_review(result, key_prefix=f"history_{selected}")
        st.download_button("Export review register CSV", reviews.to_csv(index=False), file_name="hvacr_review_register.csv", mime="text/csv")

elif page == "Company standards":
    page_header("Company standards", "Set the evidence and approval rules applied to every technical review.", "Configuration")
    st.write("Add rules that should be applied to every review, such as evidence required before compressor replacement or labour policies.")
    default_standards = st.session_state.get("company_standards", """Compressor replacement requires voltage, current, winding resistance, insulation resistance, operating pressures and likely cause of failure.
TXV replacement requires confirmation of charge, airflow/load, bulb mounting and insulation, equaliser condition, restriction checks and response to adjustment.
All labour above 8 hours must include a visit-by-visit task breakdown.
Photos should support major defects and completed repairs.""")
    standards = st.text_area("Standards and approval rules", value=default_standards, height=360)
    if st.button("Save standards", type="primary"):
        st.session_state["company_standards"] = standards
        st.success("Standards saved for this app session. For permanent cloud storage, add a hosted database in the production build.")

elif page == "Settings":
    page_header("Settings", "Configure Live AI, pilot access and production-readiness requirements.", "Administration")
    st.markdown("### Live AI")
    st.code('OPENAI_API_KEY = "sk-proj-..."\nOPENAI_MODEL = "gpt-4.1-mini"\nAPP_PASSWORD = "optional-private-pilot-password"', language="toml")
    st.info("Add these values in Streamlit → App settings → Secrets. Never store an API key in GitHub.")
    st.markdown("### Pilot status")
    st.write("This build is suitable for controlled testing with de-identified reports. Before charging customers, add proper user accounts, permanent encrypted storage, subscription billing, audit logs, privacy terms and a security review.")
